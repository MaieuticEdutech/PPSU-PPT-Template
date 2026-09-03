"""PPSU visual constants for the SLM builder — colors, fonts, shared helpers.

Colors are NOT guessed from the sample PDF's screen render; they're pulled
directly from ppsu1/template's slide master (see the extraction this was
built from) so the SLM matches the same brand chrome the PPT designer uses.
The generic theme1.xml colour scheme in that template is just Office's
unmodified default — the real PPSU palette lives in the master's own shape
fills, not the theme swatch.
"""
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

BRAND_FONT = "Calibri"      # matches ppsu1's BRAND_FONT — one PPSU voice across tools
MONO_FONT = "Consolas"      # code blocks

NAVY = RGBColor(0x0E, 0x28, 0x41)     # headings, cover dark bar
RED = RGBColor(0xD8, 0x18, 0x1F)      # SLM banner
ORANGE = RGBColor(0xF4, 0x78, 0x20)   # accent shapes
GOLD = RGBColor(0xFB, 0xB2, 0x17)     # "MSc SEMESTER (I)" style text
GREY_TEXT = RGBColor(0x5A, 0x5A, 0x5A)
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Cell shading fills (hex strings, no '#') for the boxed block types.
FILL_DID_YOU_KNOW = "E7E6E6"     # light grey — matches the sample's grey box
FILL_THINK_AND_APPLY = "FCE4D6"  # light orange tint — matches the sample's peach box
FILL_PROBLEM = "FCE4D6"          # spec: orange boxes for problem/solution
FILL_CODE = "F2F2F2"             # pale grey code background

TITLE_PT = 24
H1_PT = 20          # "N.N Section Title"
H2_PT = 15          # "N.N.N Subsection Title"
BODY_PT = 11
CAPTION_PT = 10


import brands

# the active brand profile (see brands.py). set_brand() swaps it and syncs
# the legacy module attributes the builder reads.
BRAND = brands.PPSU
_DEFAULTS = dict(NAVY=NAVY)


def set_brand(key):
    """Activate a built-in brand profile ('ppsu' | 'reva'): loads its
    values into this module's attributes. Called by build() before every
    render, so brands never leak between builds."""
    import sys
    mod = sys.modules[__name__]
    brand = brands.PROFILES[key]
    mod.BRAND = brand
    mod.BRAND_FONT = brand["font"]
    mod.TITLE_PT = brand["title_pt"]
    mod.H1_PT = brand["h1_pt"]
    mod.H2_PT = brand["h2_pt"]
    mod.BODY_PT = brand["body_pt"]
    mod.CAPTION_PT = brand["caption_pt"]
    mod.NAVY = RGBColor.from_string(brand["heading_color"])
    mod.FILL_DID_YOU_KNOW = brand["fill_did_you_know"]
    mod.FILL_THINK_AND_APPLY = brand["fill_think_apply"]
    mod.FILL_PROBLEM = brand["fill_problem"]
    mod.FILL_CODE = brand["fill_code"]


def apply_profile(profile, brand_key="ppsu"):
    """Layering order for every build: brand base theme first (reset), then
    the optional uploaded reference-PDF overrides on top. Mutates THIS
    module — the builder reads styles via module attributes."""
    import sys
    mod = sys.modules[__name__]
    set_brand(brand_key)
    if not profile:
        return
    if profile.get("font_name"):
        mod.BRAND_FONT = profile["font_name"]
    if profile.get("body_pt"):
        mod.BODY_PT = int(profile["body_pt"])
        mod.CAPTION_PT = max(8, int(profile["body_pt"]) - 1)
    if profile.get("h1_pt"):
        mod.H1_PT = int(profile["h1_pt"])
        mod.H2_PT = max(11, int(profile["h1_pt"]) - 5)
    if profile.get("heading_color"):
        try:
            mod.NAVY = RGBColor.from_string(profile["heading_color"])
        except Exception:
            pass


def set_run(run, *, size=None, bold=False, italic=False, color=None,
            font=None):
    # defaults resolve at CALL time so apply_profile() overrides are seen
    # (a def-time default would freeze the original constants)
    font = font or BRAND_FONT
    size = size if size is not None else BODY_PT
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    # east-asian font element must also be set or Word can silently fall
    # back to a different font for some code paths
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font)
    return run


def add_paragraph(doc_or_cell, text="", *, size=None, bold=False,
                   italic=False, color=None, font=None,
                   align=None, space_after=6):
    p = doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        set_run(p.add_run(text), size=size, bold=bold, italic=italic,
                color=color, font=font)
    return p


def shade_cell(cell, hex_fill):
    """Fill a table cell's background — python-docx has no high-level API
    for this, so drop to the raw <w:shd> element directly."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def remove_table_borders(table):
    """Used for the boxed-callout tables (did_you_know / think_and_apply /
    code) so they read as a shaded box, not a bordered grid."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tblPr.append(borders)


def add_box(doc, *, fill):
    """A single-cell, borderless, shaded table used as a callout box
    (did_you_know / think_and_apply / problem / code). Returns the cell —
    caller fills it with paragraphs."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    remove_table_borders(table)
    return cell


def body_para(doc_or_cell, text, *, space_after=8):
    """A body-text paragraph obeying the active brand: colour, justification
    and line spacing (REVA: #333333, justified, 1.5; PPSU: plain)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _AL
    color = (RGBColor.from_string(BRAND["body_color"])
             if BRAND.get("body_color") else None)
    p = add_paragraph(doc_or_cell, text, color=color, space_after=space_after)
    if BRAND.get("justify_body"):
        p.alignment = _AL.JUSTIFY
    if BRAND.get("line_spacing"):
        p.paragraph_format.line_spacing = BRAND["line_spacing"]
    return p


def shade_paragraph(paragraph, hex_fill):
    """Paragraph background shading (the REVA heading bars)."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    pPr.append(shd)


def paragraph_border(paragraph, edge, hex_color, size=12):
    """A single border line on a paragraph (REVA's orange header/footer
    rules). edge: 'top' or 'bottom'."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    el = OxmlElement(f'w:{edge}')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), str(size))
    el.set(qn('w:space'), '4')
    el.set(qn('w:color'), hex_color)
    pBdr.append(el)


def set_table_border_color(table, hex_color):
    """Recolour a table's grid borders (REVA: #CCCCCC)."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), hex_color)
        borders.append(el)
    tblPr.append(borders)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    run._r.addnext(fld)


def add_toc_field(paragraph):
    """A REAL, updateable Word TOC field (picks up Heading 1-3 styled
    paragraphs) — not a static hand-typed page-number list. Word shows a
    grey placeholder until the user updates fields (F9, or opens with
    'update fields on open' set, or Word prompts automatically depending
    on settings) — this is normal Word TOC-field behaviour, not a bug."""
    run = paragraph.add_run()
    r = run._r

    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')

    placeholder = OxmlElement('w:t')
    placeholder.text = "Right-click and choose \"Update Field\" to generate the table of contents."

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')

    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(placeholder)
    r.append(fld_end)
