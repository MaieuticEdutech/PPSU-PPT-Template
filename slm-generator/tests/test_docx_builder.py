#!/usr/bin/env python3
"""Golden-JSON -> docx tests for Phase 1 (see CLAUDE.md's "Tests ship with
features"). Plain assertions, no pytest — the builder has no dependency on
a test framework and this avoids adding one just to check a docx opens
correctly and has the right shape.

Run: python tests/test_docx_builder.py
"""
import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent / "backend"))

from docx import Document
from docx.oxml.ns import qn

from docx_builder import build

HERE = Path(__file__).parent
OUT = HERE.parent / "output"
OUT.mkdir(exist_ok=True)

passed = failed = 0


def check(name, condition):
    global passed, failed
    if condition:
        passed += 1
        print(f"  OK   {name}")
    else:
        failed += 1
        print(f"  FAIL {name}")


def all_text(doc):
    """doc.paragraphs only returns paragraphs that are DIRECT children of
    the body -- it silently skips anything nested inside a table cell
    (every did_you_know/think_and_apply/code/figure/problem box, PLUS the
    glossary and content tables). Walk the real XML tree instead."""
    from docx.oxml.ns import qn as _qn
    parts = []
    for t in doc.element.body.iter():
        tag = t.tag
        if tag == _qn('w:t') and t.text:
            parts.append(t.text)
        elif tag in (_qn('w:p'), _qn('w:tr')):
            parts.append("\n")
    return "".join(parts)


def heading_texts(doc, style_prefix="Heading"):
    return [p.text for p in doc.paragraphs
            if p.style and p.style.name.startswith(style_prefix)]


def has_toc_field(doc):
    body = doc.element.body
    for instr in body.iter(qn('w:instrText')):
        if instr.text and 'TOC' in instr.text:
            return True
    return False


def content_table_count(doc):
    """Real bordered content tables (Table Grid style) — excludes the
    borderless single-cell 'box' tables used for callouts/code/figures."""
    return sum(1 for t in doc.tables if t.style and t.style.name == "Table Grid")


# ---------------------------------------------------------------------------
print("=== golden_unit1.json (real reference sample, representative subset) ===")
data = json.loads((HERE / "golden_unit1.json").read_text(encoding="utf-8"))
doc = build(data)
out_path = OUT / "golden_unit1.docx"
doc.save(str(out_path))
check("file written and non-trivial size", out_path.stat().st_size > 20_000)

reopened = Document(str(out_path))
text = all_text(reopened)

check("cover banner text present", "SELF-LEARNING MATERIAL" in text)
check("unit heading present", "Unit 01: Foundations of Discrete Mathematics" in text)
check("real Word TOC field present (not a static list)", has_toc_field(reopened))
check("all 5 learning objectives' verbs present",
      all(v in text for v in ("Explain", "Identify", "Apply", "Illustrate", "Analyse")))

headings = heading_texts(reopened)
check("section headings present (1.1/1.2/1.3)",
      any("1.1 " in h for h in headings) and any("1.2 " in h for h in headings)
      and any("1.3 " in h for h in headings))
check("subsection headings present (1.1.1/1.1.2/1.1.3)",
      all(any(f"1.1.{i} " in h for h in headings) for i in (1, 2, 3)))
check("back-matter headings present (Summary/Glossary/Case Study/"
      "Self-Assessment/Terminal/Answers/References)",
      all(any(k in h for h in headings) for k in
          ("Summary", "Glossary", "Case Study", "Self-Assessment",
           "Terminal Questions", "Answers", "References")))

check("code block text preserved verbatim",
      'customers_campaign_a = {"Anu", "Raj", "Meera", "Anu"}' in text)
check("did_you_know box text preserved",
      "‘discreet’ means" in text)
check("think_and_apply box text preserved",
      "Gold’ members and ‘Frequent Buyers" in text)
check("figure placeholder + real caption both present",
      "FIGURE PLACEHOLDER" in text and "Figure 1: Discrete Mathematics" in text)

# content tables: 3 in-section tables (1.1.1, 1.2.3, 1.3.1) + glossary = 4
check(f"content table count == 4 (got {content_table_count(reopened)})",
      content_table_count(reopened) == 4)

glossary_table = next(t for t in reopened.tables
                      if t.style and t.style.name == "Table Grid"
                      and t.rows[0].cells[0].text == "Term")
glossary_terms = [r.cells[0].text for r in glossary_table.rows[1:]]
check("glossary rows are alphabetised",
      glossary_terms == sorted(glossary_terms, key=str.lower))
check("case study title present", "Cleaning Overlapping Customer Lists" in text)
check("all 3 case-study questions rendered verbatim",
      all(q in text for q in data["case_study"]["questions"]))
check("all 5 MCQs present with options", text.count(") ") >= 5 * 4)
check("MCQ answer key resolves correctly (Q1 -> B) Distinct, countable objects)",
      "1. B) Distinct, countable objects" in text)
# Matches the REAL sample's structure (checked against the PDF): 1.8 lists
# bare numbered questions, 1.9.2 lists bare numbered answers -- the question
# text is never repeated alongside its answer.
check("terminal question appears once (question-only, under N.8)",
      text.count("difference between a subset and a proper subset") == 1)
check("terminal answer appears once (answer-only, under N.9.2)",
      text.count("A proper subset contains only some elements") == 1)
check("references present", "Rosen, K. H." in text)

header_text = reopened.sections[0].header.paragraphs[0].text
footer_text = reopened.sections[0].footer.paragraphs[0].text
check("header shows unit number", "Unit 01" in header_text)
check("footer shows course code + name", "ICCS7010" in footer_text
      and "INFORMATION SECURITY AND APPLICATIONS" in footer_text)

# ---------------------------------------------------------------------------
print("\n=== block_types_synthetic.json (problem/solution + key_takeaway) ===")
syn = json.loads((HERE / "block_types_synthetic.json").read_text(encoding="utf-8"))
syn_doc = build(syn)
syn_path = OUT / "block_types_synthetic.docx"
syn_doc.save(str(syn_path))
syn_reopened = Document(str(syn_path))
syn_text = all_text(syn_reopened)

check("problem label + statement rendered", "Problem 99.1" in syn_text
      and "Compute 23 + 12 + 16 + 9" in syn_text)
check("solution rendered with verification tick",
      "23+12+16+9=60 ✓" in syn_text)
check("key_takeaway rendered with bold lead-in",
      "Key Takeaway: " in syn_text
      and "Worked problems always end with a verification step." in syn_text)

# ---------------------------------------------------------------------------
print(f"\n{passed} passed, {failed} failed")
sys.exit(1 if failed else 0)
