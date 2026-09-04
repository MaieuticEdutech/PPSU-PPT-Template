#!/usr/bin/env python3
"""docx_builder.py — renders the SLM intermediate JSON into a PPSU-branded
.docx. Deterministic: JSON in, document out, no AI calls, no randomness (see
CLAUDE.md's "Conventions" — this is Phase 1, the half that must be provable
without an LLM in the loop).

Schema note: this deliberately departs from the JSON sketch in CLAUDE.md in
a few places, based on what the REAL reference sample (P P Savani University
ICCS7010 Unit 1, supplied 2026-09-02) actually contains rather than what the
spec assumed before any real sample existed:

  - `code` is a new block type. The spec's block list (prose/table/
    did_you_know/problem+solution/key_takeaway/think_and_apply/figure) has
    no slot for the Python/SQL code snippets that appear constantly in this
    (data-science-flavoured) unit -- 1.1.3, 1.2.1, 1.2.2, 1.2.3, 1.2.4,
    1.3.2 (x2), 1.3.3 all have one. Rendered as a shaded, monospace box.
  - `problem`/`solution` and `key_takeaway` are implemented (the spec
    describes them and a different, more traditionally mathematical unit
    may well use them) but the real sample uses NEITHER -- it's an
    "algorithm + code example" unit, not a "worked numeric problem" one.
    Block-type usage is topic-dependent; the builder must render whichever
    subset of the 8 known types a given unit's JSON actually contains.
  - `case_study.questions` is a flat list of question strings, not the
    spec's `[{"q", "a"}]` pairs -- the real sample's case-study questions
    have no provided answers at all (they're left open, presumably for
    classroom discussion or SME-authored answers later).
  - Answers are NOT duplicated into a separate top-level "answers" object.
    Each self-assessment MCQ / terminal question already carries its own
    correct answer (`answer`, optionally `why`); the builder renders the
    bare question under N.7/N.8 and the answer view under N.9 by reading
    the SAME objects twice, rather than requiring the same content be
    written out in two places in the JSON.
  - `self_assessment.fill_blanks` may be an empty list (the real sample has
    no "B. Fill in the Blanks" subsection at all) -- the builder skips that
    heading entirely when the list is empty, rather than rendering an empty
    section.

See tests/golden_unit1.json for a representative (not exhaustive -- the real
unit is 33 pages) transcription proving every block type renders, plus one
clearly-marked synthetic subsection covering problem/key_takeaway, which the
real sample doesn't exercise.
"""
import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import brands
import styles as st
from styles import (add_box, add_page_number_field, add_paragraph,
                    add_toc_field, body_para, paragraph_border,
                    set_run, set_table_border_color, shade_cell,
                    shade_paragraph)


def _try_logo(target, logo, height):
    """Embed the logo into a run/document, never letting an unreadable
    image kill the whole build (upload-time validation should prevent it,
    but a decorative image is never worth failing a unit over). Returns
    True when embedded."""
    try:
        target.add_picture(str(logo), height=height)
        return True
    except Exception:
        return False


def _brand_caption(caption):
    """Captions are generated as 'Figure N: ...' — REVA's convention is
    'Fig. N: ...' (style guide caption examples). Transformed at render
    time so the SAME unit JSON can target either brand."""
    prefix = st.BRAND["fig_prefix"]
    if prefix != "Figure" and caption.startswith("Figure "):
        return prefix + caption[len("Figure"):]
    return caption


# ---------------------------------------------------------------------------
# cover page + running headers/footers
# ---------------------------------------------------------------------------

def _build_cover(doc, meta, logo=None):
    if logo and _try_logo(doc, logo, Inches(0.7)):
        doc.add_paragraph()
    for _ in range(2 if logo else 3):
        doc.add_paragraph()

    p = add_paragraph(doc, meta.get("course_code", ""), size=12, bold=True,
                       color=st.NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, meta.get("programme", ""), size=16, bold=True,
                  color=st.NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, meta.get("course_name", ""), size=16, bold=True,
                  color=st.NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)

    banner_cell = add_box(doc, fill="D8181F")
    add_paragraph(banner_cell, "SELF-LEARNING MATERIAL", size=14, bold=True,
                  color=st.WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    for _ in range(2):
        doc.add_paragraph()

    add_paragraph(doc, meta.get("programme", ""), size=28, bold=True,
                  color=st.GOLD, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, meta.get("course_name", ""), size=20, bold=True,
                  color=RGBColor(0, 0, 0), align=WD_ALIGN_PARAGRAPH.CENTER)

    for _ in range(4):
        doc.add_paragraph()

    # Cover photography is a DTP asset this builder does not have; a
    # labelled placeholder box stands in so the layout is provable and the
    # DTP team knows exactly what to drop in.
    photo_cell = add_box(doc, fill="E7E6E6")
    add_paragraph(photo_cell, "[ COVER PHOTOGRAPHY PLACEHOLDER — DTP team to "
                  "insert PPSU stock imagery ]", size=10, italic=True,
                  color=st.GREY_TEXT, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()


def _add_running_header_footer(doc, meta, logo=None):
    section = doc.sections[0]
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    brand = st.BRAND
    unit_no = meta.get("unit_number", "")
    n_label = f"{unit_no:02d}" if isinstance(unit_no, int) else str(unit_no)
    if brand.get("header_text") == "unit_name":
        # REVA header content: "Unit Name | Unit Number" (+ logo right)
        unit_text = f'{meta.get("unit_title", "")} | Unit {n_label}'
    else:
        unit_text = f"Unit {n_label}"

    from docx.enum.text import WD_TAB_ALIGNMENT
    hp = section.header.paragraphs[0]
    usable = (section.page_width - section.left_margin
              - section.right_margin)
    hf_color = brand.get("hf_border_color")
    if logo and brand.get("header_layout") == "logo_right":
        # REVA: text left, logo right
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.tab_stops.add_tab_stop(usable,
                                                   WD_TAB_ALIGNMENT.RIGHT)
        set_run(hp.add_run(unit_text), size=11, bold=True, color=st.NAVY)
        _try_logo(hp.add_run("\t"), logo, Inches(0.32))
    elif logo:
        # PPSU: logo left, unit label right
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.tab_stops.add_tab_stop(usable,
                                                   WD_TAB_ALIGNMENT.RIGHT)
        _try_logo(hp.add_run(), logo, Inches(0.32))
        hp.add_run("\t")
        set_run(hp.add_run(unit_text), size=11, bold=True, color=st.NAVY)
    else:
        hp.alignment = (WD_ALIGN_PARAGRAPH.LEFT
                        if brand.get("header_text") == "unit_name"
                        else WD_ALIGN_PARAGRAPH.RIGHT)
        set_run(hp.add_run(unit_text), size=11, bold=True, color=st.NAVY)
    if hf_color:                          # REVA: orange rule under header
        paragraph_border(hp, "bottom", hf_color)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    code = meta.get("course_code", "")
    name = meta.get("course_name", "")
    footer_color = RGBColor.from_string(brand.get("footer_color", "0E2841"))
    if brand.get("header_text") == "unit_name":
        # REVA footer content: "Course Code | Course Name" + page (right)
        footer_text = f"{code} | {name}" if code or name else ""
    else:
        footer_text = f"{code}: {name.upper()}" if code or name else ""
    set_run(fp.add_run(footer_text), size=brand.get("footer_pt", 9),
            bold=brand.get("footer_bold", True), color=footer_color)
    fp.add_run("\t\t")
    add_page_number_field(fp)
    if hf_color:                          # REVA: orange rule above footer
        paragraph_border(fp, "top", hf_color)


# ---------------------------------------------------------------------------
# unit heading, TOC, introduction, learning objectives
# ---------------------------------------------------------------------------

def _build_unit_heading_and_toc(doc, meta):
    h = doc.add_heading(level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    unit_no = meta.get("unit_number", "")
    unit_label = f"{unit_no:02d}" if isinstance(unit_no, int) else str(unit_no)
    set_run(h.add_run(f"Unit {unit_label}: {meta.get('unit_title', '')}"),
            size=st.TITLE_PT, bold=True, color=st.NAVY)

    toc_title = add_paragraph(doc, "Table of Contents", size=14, bold=True,
                              color=st.NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    toc_para = doc.add_paragraph()
    add_toc_field(toc_para)
    doc.add_page_break()


def _build_introduction(doc, intro_paragraphs):
    _heading(doc, "Introduction", level=1)
    for para in intro_paragraphs:
        body_para(doc, para)


def _build_learning_objectives(doc, objectives):
    _heading(doc, "Learning Objectives", level=1)
    add_paragraph(doc, "By the end of this unit, you will be able to:",
                  size=st.BODY_PT, space_after=6)
    for obj in objectives:
        p = doc.add_paragraph(style="List Bullet")
        set_run(p.add_run(obj.get("verb", "")), bold=True, size=st.BODY_PT)
        set_run(p.add_run(" " + obj.get("rest", "")), size=st.BODY_PT)


# ---------------------------------------------------------------------------
# heading helpers (drive the real Word TOC — must be built-in Heading styles)
# ---------------------------------------------------------------------------

def _heading(doc, text, level):
    h = doc.add_heading(level=level)
    brand = st.BRAND
    size = {1: st.H1_PT, 2: st.H2_PT,
            3: brand.get("h3_pt", st.H2_PT - 2)}.get(level, st.H2_PT - 2)
    if brand.get("heading_upper"):        # REVA: Bold, Uppercase headings
        text = text.upper()
    fill = brand.get("h_fills", {}).get(level)
    text_hex = brand.get("h_text_colors", {}).get(level)
    color = (RGBColor.from_string(text_hex) if text_hex else st.NAVY)
    set_run(h.add_run(text), size=size, bold=True, color=color)
    if fill:                               # REVA: coloured heading bars
        shade_paragraph(h, fill)
    return h


# ---------------------------------------------------------------------------
# content blocks
# ---------------------------------------------------------------------------

def _render_table(doc, block):
    brand = st.BRAND
    if block.get("caption"):
        add_paragraph(doc, block["caption"], size=st.CAPTION_PT, bold=True,
                      align=(WD_ALIGN_PARAGRAPH.CENTER
                             if brand.get("caption_center") else None),
                      space_after=4)
    columns = block.get("columns", [])
    rows = block.get("rows", [])
    table = doc.add_table(rows=1 + len(rows), cols=max(len(columns), 1))
    table.style = "Table Grid"
    if brand.get("table_border"):          # REVA: #CCCCCC grid
        set_table_border_color(table, brand["table_border"])
    for i, col in enumerate(columns):
        cell = table.rows[0].cells[i]
        shade_cell(cell, brand["table_header_fill"])
        p = cell.paragraphs[0]
        set_run(p.add_run(col), bold=True, size=st.BODY_PT - 1,
                color=RGBColor.from_string(brand["table_header_text"]))
    body_color = (RGBColor.from_string(brand["body_color"])
                  if brand.get("body_color") else None)
    for r, row in enumerate(rows, start=1):
        # REVA: alternate row shading (#F2F2F2 on even data rows)
        if brand.get("table_alt_row") and r % 2 == 0:
            for cell in table.rows[r].cells:
                shade_cell(cell, brand["table_alt_row"])
        for c, val in enumerate(row):
            if c >= len(table.rows[r].cells):
                continue
            p = table.rows[r].cells[c].paragraphs[0]
            set_run(p.add_run(str(val)), size=st.BODY_PT - 1,
                    color=body_color)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _render_box(doc, text, *, fill, label=None, label_color=None):
    if label_color is None:
        # boxes carry body-coloured text where the brand defines one
        # (REVA: #333333 on the yellow/green fills), else the heading tone
        label_color = (RGBColor.from_string(st.BRAND["body_color"])
                       if st.BRAND.get("body_color") else st.NAVY)
    cell = add_box(doc, fill=fill)
    if label:
        add_paragraph(cell, label, size=st.BODY_PT, bold=True,
                      color=label_color, space_after=4)
    body_para(cell, text, space_after=2)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _render_code(doc, text):
    cell = add_box(doc, fill=st.FILL_CODE)
    lines = text.split("\n")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        set_run(p.add_run(line if line else " "), size=st.BODY_PT - 1,
                font=st.MONO_FONT)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _render_figure(doc, block):
    # a generated diagram (block["image"], from figure_render) is embedded
    # when present and readable; anything wrong with it falls back to the
    # DTP placeholder box — same belt as _try_logo
    embedded = False
    image = block.get("image")
    if image and Path(image).is_file():
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(image), width=Inches(5.5))
            embedded = True
        except Exception:                               # noqa: BLE001
            pass
    if not embedded:
        cell = add_box(doc, fill="F2F2F2")
        add_paragraph(cell,
                      "[ FIGURE PLACEHOLDER — DTP team to insert artwork ]",
                      size=st.BODY_PT, italic=True, color=st.GREY_TEXT,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    if block.get("caption"):
        add_paragraph(doc, _brand_caption(block["caption"]),
                      size=st.CAPTION_PT, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)


def _render_block(doc, block):
    t = block.get("type")
    if t == "prose":
        body_para(doc, block.get("text", ""))
    elif t == "table":
        _render_table(doc, block)
    elif t == "did_you_know":
        # brand-labelled aside: "Did you know?" (PPSU) / "Study Note" (REVA)
        _render_box(doc, block.get("text", ""), fill=st.FILL_DID_YOU_KNOW,
                    label=st.BRAND["did_you_know_label"])
    elif t == "code":
        _render_code(doc, block.get("text", ""))
    elif t == "problem":
        _render_box(doc, block.get("statement", ""), fill=st.FILL_PROBLEM,
                    label=block.get("label", "Problem"))
        _render_box(doc, block.get("solution", ""), fill=st.FILL_PROBLEM,
                    label="Solution")
    elif t == "key_takeaway":
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        set_run(p.add_run("Key Takeaway: "), bold=True, size=st.BODY_PT)
        set_run(p.add_run(block.get("text", "")), size=st.BODY_PT)
    elif t == "think_and_apply":
        # "Think and Apply" (PPSU) / "Activity" (REVA)
        _render_box(doc, block.get("text", ""), fill=st.FILL_THINK_AND_APPLY,
                    label=st.BRAND["think_apply_label"])
    elif t == "figure":
        _render_figure(doc, block)
    else:
        raise ValueError(f"unknown block type: {t!r}")


# ---------------------------------------------------------------------------
# sections / subsections
# ---------------------------------------------------------------------------

def _build_sections(doc, sections):
    for sec in sections:
        _heading(doc, f'{sec["number"]} {sec["title"]}', level=1)
        if sec.get("intro"):
            body_para(doc, sec["intro"])
        for sub in sec.get("subsections", []):
            _heading(doc, f'{sub["number"]} {sub["title"]}', level=2)
            for block in sub.get("blocks", []):
                _render_block(doc, block)


# ---------------------------------------------------------------------------
# back matter
# ---------------------------------------------------------------------------

def _boxed(doc, fill):
    """Brand-conditional container: a shaded box cell when the brand
    defines a fill for this element (REVA's summary/case/terminal boxes),
    else the document itself."""
    return add_box(doc, fill=fill) if fill else doc


def _build_summary(doc, summary_number, bullets):
    _heading(doc, f"{summary_number} Summary", level=1)
    target = _boxed(doc, st.BRAND.get("summary_box"))
    for b in bullets:
        p = target.add_paragraph(style="List Bullet")
        set_run(p.add_run(b), size=st.BODY_PT)


def _build_glossary(doc, number, terms):
    _heading(doc, f"{number} Glossary", level=1)
    brand = st.BRAND
    table = doc.add_table(rows=1 + len(terms), cols=2)
    table.style = "Table Grid"
    if brand.get("table_border"):
        set_table_border_color(table, brand["table_border"])
    for i, hdr in enumerate(("Term", "Definition")):
        cell = table.rows[0].cells[i]
        shade_cell(cell, brand["table_header_fill"])
        set_run(cell.paragraphs[0].add_run(hdr), bold=True,
                color=RGBColor.from_string(brand["table_header_text"]),
                size=st.BODY_PT - 1)
    for r, entry in enumerate(sorted(terms, key=lambda e: e["term"].lower()),
                              start=1):
        if brand.get("glossary_fill"):     # REVA: light-grey glossary rows
            for cell in table.rows[r].cells:
                shade_cell(cell, brand["glossary_fill"])
        set_run(table.rows[r].cells[0].paragraphs[0].add_run(entry["term"]),
                bold=True, size=st.BODY_PT - 1)
        set_run(table.rows[r].cells[1].paragraphs[0]
                .add_run(entry["definition"]), size=st.BODY_PT - 1)


def _build_case_study(doc, number, case):
    _heading(doc, f"{number} Case Study", level=1)
    target = _boxed(doc, st.BRAND.get("case_box"))
    if case.get("title"):
        add_paragraph(target, case["title"], size=st.BODY_PT + 1, bold=True,
                      space_after=6)
    bg = case.get("background", "")
    paras = bg if isinstance(bg, list) else [bg]
    for para in paras:
        if para:
            body_para(target, para)
    if case.get("questions"):
        add_paragraph(target, "Case Study Questions:", size=st.BODY_PT,
                      bold=True, space_after=4)
        for q in case["questions"]:
            p = target.add_paragraph(style="List Bullet")
            set_run(p.add_run(q), size=st.BODY_PT)


def _build_self_assessment(doc, number, sa):
    _heading(doc, f"{number} Self-Assessment Questions", level=1)
    mcqs = sa.get("mcq", [])
    if mcqs:
        add_paragraph(doc, "A. Multiple Choice Questions", size=st.BODY_PT,
                      bold=True, space_after=6)
        letters = "abcd"
        for i, item in enumerate(mcqs, start=1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            set_run(p.add_run(f'{i}. {item["q"]}'), bold=True, size=st.BODY_PT)
            for letter, opt in zip(letters, item.get("options", [])):
                op = doc.add_paragraph()
                op.paragraph_format.left_indent = Inches(0.3)
                op.paragraph_format.space_after = Pt(0)
                set_run(op.add_run(f"{letter}) {opt}"), size=st.BODY_PT)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    blanks = sa.get("fill_blanks", [])
    if blanks:
        add_paragraph(doc, "B. Fill in the Blanks", size=st.BODY_PT, bold=True,
                      space_after=6)
        for i, item in enumerate(blanks, start=1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            set_run(p.add_run(f'{i}. {item["q"]}'), size=st.BODY_PT)


def _build_terminal(doc, number, terminal):
    _heading(doc, f"{number} Terminal Questions", level=1)
    target = _boxed(doc, st.BRAND.get("terminal_box"))
    for label, key in (("Short Questions", "short"), ("Long Questions", "long")):
        items = terminal.get(key, [])
        if not items:
            continue
        add_paragraph(target, label, size=st.BODY_PT, bold=True,
                      space_after=6)
        for i, item in enumerate(items, start=1):
            p = target.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            set_run(p.add_run(f'{i}. {item["q"]}'), size=st.BODY_PT)


def _mcq_answer_line(item):
    letters = "abcd"
    idx = letters.index(item["answer"].lower())
    text = item.get("options", [None] * 4)[idx]
    line = f'{item["answer"].upper()}) {text}'
    if item.get("why"):
        line += f' — {item["why"]}'
    return line


def _build_answers(doc, number, sa, terminal):
    _heading(doc, f"{number} Answers", level=1)

    mcqs = sa.get("mcq", [])
    blanks = sa.get("fill_blanks", [])
    if mcqs or blanks:
        _heading(doc, f"{number}.1 Self-Assessment Answers", level=2)
        if mcqs:
            add_paragraph(doc, "A. Multiple Choice Questions", size=st.BODY_PT,
                          bold=True, space_after=4)
            for i, item in enumerate(mcqs, start=1):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                set_run(p.add_run(f'{i}. {_mcq_answer_line(item)}'),
                        size=st.BODY_PT)
        if blanks:
            add_paragraph(doc, "B. Fill in the Blanks", size=st.BODY_PT,
                          bold=True, space_after=4)
            for i, item in enumerate(blanks, start=1):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                set_run(p.add_run(f'{i}. {item["answer"]}'), size=st.BODY_PT)

    short_items = terminal.get("short", [])
    long_items = terminal.get("long", [])
    if short_items or long_items:
        _heading(doc, f"{number}.2 Terminal Answers", level=2)
        for label, items in (("Short Answers", short_items),
                             ("Long Answers", long_items)):
            if not items:
                continue
            add_paragraph(doc, label, size=st.BODY_PT, bold=True, space_after=4)
            for i, item in enumerate(items, start=1):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(8)
                set_run(p.add_run(f'{i}. '), bold=True, size=st.BODY_PT)
                set_run(p.add_run(item.get("answer", "")), size=st.BODY_PT)


def _build_references(doc, number, refs):
    # "References" (PPSU) / "Suggested Books and References" (REVA structure)
    _heading(doc, f'{number} {st.BRAND["references_title"]}', level=1)
    for ref in refs:
        p = doc.add_paragraph(style="List Bullet")
        set_run(p.add_run(ref), size=st.BODY_PT)


# ---------------------------------------------------------------------------
# entry point
# ---------------------------------------------------------------------------

def _apply_page_geometry(doc, profile):
    """Page size + margins from the reference-PDF profile (mm), when set."""
    if not profile or not profile.get("page_width_mm"):
        return
    from docx.shared import Mm
    section = doc.sections[0]
    section.page_width = Mm(profile["page_width_mm"])
    section.page_height = Mm(profile["page_height_mm"])
    for attr, key in (("left_margin", "margin_left_mm"),
                      ("right_margin", "margin_right_mm"),
                      ("top_margin", "margin_top_mm"),
                      ("bottom_margin", "margin_bottom_mm")):
        val = profile.get(key)
        # a reference PDF's text can start absurdly close to the edge
        # (headers/footers count as text) — clamp to sane print margins
        if val and 8 <= val <= 40:
            setattr(section, attr, Mm(val))


def build(data: dict, *, brand="ppsu", use_branding=True) -> Document:
    """Render a unit for a brand ('ppsu' | 'reva', see brands.py).
    Layering: brand base theme, then uploaded reference-PDF style profile,
    then uploaded logo. use_branding=False skips the uploaded assets (used
    by tests asserting a brand's base look)."""
    logo = profile = None
    if use_branding:
        import branding
        logo = branding.logo_path()
        profile = branding.load_profile()
    st.apply_profile(profile, brand_key=brand)   # resets + layers
    doc = Document()
    _apply_page_geometry(doc, profile)
    _add_running_header_footer(doc, data.get("meta", {}), logo=logo)
    _build_cover(doc, data.get("meta", {}), logo=logo)
    _build_unit_heading_and_toc(doc, data.get("meta", {}))
    _build_introduction(doc, data.get("introduction", []))
    _build_learning_objectives(doc, data.get("learning_objectives", []))
    _build_sections(doc, data.get("sections", []))

    unit_no = data.get("meta", {}).get("unit_number", 1)
    n = str(unit_no)
    _build_summary(doc, f"{n}.4", data.get("summary", []))
    _build_glossary(doc, f"{n}.5", data.get("glossary", []))
    _build_case_study(doc, f"{n}.6", data.get("case_study", {}))
    _build_self_assessment(doc, f"{n}.7", data.get("self_assessment", {}))
    _build_terminal(doc, f"{n}.8", data.get("terminal", {}))
    _build_answers(doc, f"{n}.9", data.get("self_assessment", {}),
                   data.get("terminal", {}))
    _build_references(doc, f"{n}.10", data.get("references", []))
    return doc


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--json", required=True, type=Path)
    ap.add_argument("--out", required=True, type=Path)
    args = ap.parse_args()
    data = json.loads(args.json.read_text(encoding="utf-8"))
    doc = build(data)
    doc.save(str(args.out))
    print(f"wrote {args.out}")


if __name__ == "__main__":
    main()
